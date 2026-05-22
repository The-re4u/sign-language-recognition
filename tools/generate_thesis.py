# coding:utf-8
"""Generate graduation thesis .docx from project data and documentation."""
import sys, os, json, datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("pip install python-docx")
    sys.exit(1)

OUTPUT = 'docs/毕业论文_章缪琪_手语实时识别系统.docx'

doc = Document()

# ========== STYLES ==========
def setup_styles():
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    pf = style.paragraph_format
    pf.line_spacing = 1.25
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def heading1(text):
    """Chapter title: Times 16pt bold, 24pt spacing"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.bold = True
    pf = p.paragraph_format
    pf.space_before = Pt(24)
    pf.space_after = Pt(24)
    return p

def heading2(text):
    """Section: Times 14pt bold, 12pt spacing"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(12)
    return p

def heading3(text):
    """Subsection: Times 12pt bold, 12pt spacing"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(12)
    return p

def body(text):
    """Body text: Times 12pt, 1.25 spacing, justified, first-line indent"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0.75)
    pf.line_spacing = 1.25
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def body_no_indent(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    pf = p.paragraph_format
    pf.line_spacing = 1.25
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def fig_placeholder(caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'[Figure — {caption}]')
    r.font.size = Pt(10); r.font.color.rgb = RGBColor(150,150,150); r.italic = True
    return p

def bullet(text):
    p = doc.add_paragraph()
    p.style = doc.styles['List Bullet']
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def newpage():
    doc.add_page_break()

def add_table(headers, rows):
    """Add formatted table"""
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9); r.font.name = 'Times New Roman'; r.bold = True
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9); r.font.name = 'Times New Roman'
    doc.add_paragraph()
    return table

setup_styles()

# ================================================================
# COVER PAGE
# ================================================================
for _ in range(5):
    doc.add_paragraph()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('本科毕业设计说明书（论文）'); r.font.name = '黑体'; r.font.size = Pt(22); r.bold = True

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Undergraduate Graduation Project Report (Thesis)'); r.font.name = 'Times New Roman'; r.font.size = Pt(11)

for _ in range(3):
    doc.add_paragraph()

# Chinese title
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('基于轻量化多模态的手语实时识别系统\n设计与实现'); r.font.name = '黑体'; r.font.size = Pt(18); r.bold = True

doc.add_paragraph()

# English title
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('DESIGN AND IMPLEMENTATION OF A LIGHTWEIGHT MULTIMODAL\nREAL-TIME SIGN LANGUAGE RECOGNITION SYSTEM')
r.font.name = 'Times New Roman'; r.font.size = Pt(14); r.bold = True

for _ in range(4):
    doc.add_paragraph()

info = [
    ('学   院：', '计算机科学与技术学院、软件学院'),
    ('专   业：', '软件工程（中外合作办学）'),
    ('班   级：', '2022 软件工程（中外合作办学）'),
    ('学   号：', '202203340226'),
    ('学生姓名：', '章缪琪'),
    ('指导老师：', '陈波'),
    ('提交日期：', '2026年6月'),
]
for label, value in info:
    p = doc.add_paragraph()
    r = p.add_run(label + value); r.font.name = '宋体'; r.font.size = Pt(14)
    p.paragraph_format.left_indent = Cm(6)

newpage()
print('[1/6] Cover page', flush=True)

# ABSTRACT
import sys; sys.stdout.flush()
print('  Starting abstract...', flush=True)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('ABSTRACT'); r.font.name = 'Times New Roman'; r.font.size = Pt(16); r.bold = True
p.paragraph_format.space_before = Pt(24); p.paragraph_format.space_after = Pt(24)

texts = json.load(open('tools/thesis_texts.json', 'r', encoding='utf-8'))
abstract_en = texts['abstract_en']
for para in abstract_en.split('\n\n'):
    if para.strip():
        body(para.strip())
